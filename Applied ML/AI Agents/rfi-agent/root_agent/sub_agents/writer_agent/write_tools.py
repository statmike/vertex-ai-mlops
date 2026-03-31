import os
from docx import Document
from openpyxl import load_workbook
from google.genai import Client
from typing import List, Dict, Any
import json
from root_agent.models.schema import RFIState, Question

def write_docx(state: RFIState, output_path: str):
    """
    Opens the original DOCX file, writes the answers back to the tracked indices,
    and saves to a new output path.
    """
    doc = Document(state.document.path)
    
    for q in state.questions:
        if q.status == "critiqued" and q.answer.text:
            loc = q.location
            if loc.type == "docx_paragraph":
                p_index = loc.paragraph_index
                if p_index is not None and p_index < len(doc.paragraphs):
                    # We append the answer directly to the question paragraph to avoid formatting breaks
                    doc.paragraphs[p_index].add_run("\n\nAnswer: " + q.answer.text).bold = True
                    
            elif loc.type == "docx_table":
                # Look up the exact table cell
                try:
                    table = doc.tables[loc.table_index]
                    row = table.rows[loc.row_index]
                    
                    # 1. Try to write to the cell *next* to it if it exists (usually the blank "Vendor Response" block)
                    if loc.cell_index + 1 < len(row.cells):
                        target_cell = row.cells[loc.cell_index + 1]
                        target_cell.text = q.answer.text
                    # 2. Check if the NEXT ROW is an explicitly blank answer block
                    elif loc.row_index + 1 < len(table.rows):
                        next_row = table.rows[loc.row_index + 1]
                        if loc.cell_index < len(next_row.cells):
                            target_cell = next_row.cells[loc.cell_index]
                            # If cell is empty or just contains a placeholder dot, write answer there
                            if len(target_cell.text.strip()) < 5:
                                target_cell.text = q.answer.text
                            else:
                                row.cells[loc.cell_index].text += "\n\nAnswer: " + q.answer.text
                        else:
                            row.cells[loc.cell_index].text += "\n\nAnswer: " + q.answer.text
                    else:
                        # 3. Otherwise, append to the question cell
                        row.cells[loc.cell_index].text += "\n\nAnswer: " + q.answer.text
                except IndexError:
                    pass # Safety catch for weird document mutations
                    
    doc.save(output_path)


def write_excel(state: RFIState, output_path: str):
    """
    Opens the original XLSX file, tracks down the exact Sheet and cell reference,
    and places the answer in the adjacent column.
    """
    wb = load_workbook(state.document.path)
    
    for q in state.questions:
        if q.status == "critiqued" and q.answer.text:
            loc = q.location
            if loc.type == "excel_cell":
                if loc.sheet_name in wb.sheetnames:
                    sheet = wb[loc.sheet_name]
                    target_cell = sheet[loc.cell_reference]
                    
                    # Assume answers go in the column immediately to the right
                    # Incrementing the column index:
                    answer_col = target_cell.column + 1
                    sheet.cell(row=target_cell.row, column=answer_col).value = q.answer.text
                    
    wb.save(output_path)


def write_pdf_fallback(state: RFIState, output_path: str):
    """
    Generates a Markdown Report of the RFI answers since we cannot write back into a PDF easily.
    Saves to the output path (with .md extension if it was .pdf).
    """
    if output_path.endswith(".pdf"):
        output_path = output_path[:-4] + ".md"
        
    with open(output_path, "w") as f:
        f.write(f"# RFI Response Report: {state.document.filename}\n\n")
        f.write(f"**Status**: {state.document.status}\n")
        f.write(f"**Total Questions**: {len(state.questions)}\n\n")
        
        for q in state.questions:
            f.write(f"## Question: {q.text}\n")
            f.write(f"- **ID**: `{q.id}`\n")
            if hasattr(q.location, 'page_number'):
                f.write(f"- **Page**: {q.location.page_number}\n")
            if q.qualification and q.qualification.question_type:
                f.write(f"- **Type**: {q.qualification.question_type}\n")
            f.write(f"- **Status**: {q.status}\n\n")
            f.write(f"### Answer\n\n{q.answer.text if q.answer.text else 'N/A'}\n\n")
            if q.answer.confidence_score:
                f.write(f"- **Confidence Score**: {int(q.answer.confidence_score * 100)}%\n")
            if q.answer.sources:
                f.write(f"- **Sources**:\n")
                for s in q.answer.sources:
                    f.write(f"  - {s}\n")
            f.write("\n---\n\n")
