import os
import time
import json
from docx import Document
from openpyxl import load_workbook
from google import genai
from pydantic import BaseModel
from typing import List, Optional

# We will use the Google GenAI SDK to parse the layout intelligently
client = genai.Client()


def call_llm_with_retry(client, *args, **kwargs):
    max_retries = 3
    base_delay = 2
    for attempt in range(max_retries):
        try:
            return client.models.generate_content(*args, **kwargs)
        except Exception as e:
            # Catch standard 429 ResourceExhausted or quota errors
            if attempt < max_retries - 1 and ("429" in str(e) or "quota" in str(e).lower() or "ResourceExhausted" in str(e)):
                # Exponential backoff: 2, 4, 8 seconds
                delay = base_delay ** (attempt + 1)
                time.sleep(delay)
            else:
                raise e
    raise Exception("Max retries exceeded for rate limits.")

class ExtractedItem(BaseModel):
    index_id: str
    question_text: str

class ExtractedQuestionsList(BaseModel):
    questions: List[ExtractedItem]

def parse_docx(file_path: str) -> List[dict]:
    """
    Parses a DOCX file, pairing index/location metadata with the text.
    Sends batches to an LLM to identify which ones are RFI questions.
    """
    doc = Document(file_path)
    
    # 1. Gather all potential text blocks
    potential_blocks = []
    
    # Process Paragraphs
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            potential_blocks.append({"loc_type": "docx_paragraph", "index": i, "text": p.text.strip(), "table_index": None, "row_index": None, "cell_index": None})
            
    # Process Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    potential_blocks.append({"loc_type": "docx_table", "table_index": t_idx, "row_index": r_idx, "cell_index": c_idx, "text": cell.text.strip(), "index": None})

    # 2. Ask LLM to filter and extract actual questions
    # 2. Ask LLM to filter and extract actual questions
    from concurrent.futures import ThreadPoolExecutor
    
    def process_docx_chunk(chunk):
        prompt = "Review the following indexed text blocks from an RFI document. Identify which ones act as standalone Questionnaire/RFI Questions that need answering. Return a list of their exact index_ids and the question text.\n\n"
        for block in chunk:
            index_id = f"p_{block['index']}" if block['index'] is not None else f"t_{block['table_index']}_r_{block['row_index']}_c_{block['cell_index']}"
            prompt += f"ID: {index_id} | Text: {block['text']}\n"
            
        response = call_llm_with_retry(
            client,
            model='gemini-2.5-pro',
            contents=prompt,
            config={
                'response_mime_type': 'application/json',
                'response_schema': ExtractedQuestionsList,
            },
        )
        
        if response.parsed and hasattr(response.parsed, 'questions'):
            return response.parsed.questions
        return []

    chunks = [potential_blocks[i:i + 50] for i in range(0, len(potential_blocks), 50)]
    all_questions = []
    
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = [executor.submit(process_docx_chunk, chunk) for chunk in chunks]
        for future in futures:
            all_questions.extend(future.result())
            
    # 3. Map back to our exact location schema
    extracted_questions = []
    for q in all_questions:
        loc = {}
        if q.index_id.startswith('p_'):
            loc = {"type": "docx_paragraph", "paragraph_index": int(q.index_id.split('_')[1])}
        elif q.index_id.startswith('t_'):
            parts = q.index_id.split('_')
            loc = {
                "type": "docx_table", 
                "table_index": int(parts[1]),
                "row_index": int(parts[3]),
                "cell_index": int(parts[5])
            }
        
        extracted_questions.append({
            "id": q.index_id,
            "text": q.question_text,
            "location": loc
        })
        
    return extracted_questions


def parse_excel(file_path: str) -> List[dict]:
    """
    Parses an Excel file to find questions.
    """
    wb = load_workbook(file_path, data_only=True)
    potential_blocks = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and cell.value.strip():
                    potential_blocks.append({
                        "loc_type": "excel_cell",
                        "sheet_name": sheet_name,
                        "cell_reference": cell.coordinate,
                        "text": cell.value.strip()
                    })
                    
    # 2. Use LLM to filter actual questions
    from concurrent.futures import ThreadPoolExecutor
    
    def process_excel_chunk(chunk):
        prompt = "Review the following indexed cells from an RFI Excel spreadsheet. Identify which ones act as the core Questionnaire/RFI questions that a vendor needs to answer. Return their index_ids and text.\n\n"
        for block in chunk:
            index_id = f"{block['sheet_name']}||{block['cell_reference']}"
            prompt += f"ID: {index_id} | Text: {block['text']}\n"
            
        response = call_llm_with_retry(
            client,
            model='gemini-2.5-pro',
            contents=prompt,
            config={
                'response_mime_type': 'application/json',
                'response_schema': ExtractedQuestionsList,
            },
        )
        
        if response.parsed and hasattr(response.parsed, 'questions'):
            return response.parsed.questions
        return []

    chunks = [potential_blocks[i:i + 50] for i in range(0, len(potential_blocks), 50)]
    all_questions = []
    
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = [executor.submit(process_excel_chunk, chunk) for chunk in chunks]
        for future in futures:
            all_questions.extend(future.result())
            
    # 3. Map back to our exact location schema
    extracted_questions = []
    for q in all_questions:
        parts = q.index_id.split('||')
        if len(parts) == 2:
            loc = {
                "type": "excel_cell",
                "sheet_name": parts[0],
                "cell_reference": parts[1]
            }
            extracted_questions.append({
                "id": q.index_id.replace("||", "_"),
                "text": q.question_text,
                "location": loc
            })
            
    return extracted_questions


def parse_pdf(file_path: str) -> List[dict]:
    """
    Parses a PDF file page by page with a sliding context window (N-1, N, N+1)
    to avoid truncation. Uses Gemini's native capability to extract questions.
    """
    from pypdf import PdfReader, PdfWriter
    from google.genai import types
    import os
    
    reader = PdfReader(file_path)
    total_pages = len(reader.pages)
    
    extracted_questions = []
    
    abs_file_path = os.path.abspath(file_path)
    temp_dir = os.path.normpath(os.path.join(os.path.dirname(abs_file_path), "../../data/temp_chunks"))
    os.makedirs(temp_dir, exist_ok=True)
    
    from concurrent.futures import ThreadPoolExecutor
    
    def process_page(i):
        writer_local = PdfWriter()
        window_pages = []
        
        if i > 1:
            writer_local.add_page(reader.pages[i-2])
            window_pages.append(i-1)
        writer_local.add_page(reader.pages[i-1]) # Target page
        window_pages.append(i)
        if i < total_pages:
            writer_local.add_page(reader.pages[i])
            window_pages.append(i+1)
            
        chunk_file = os.path.join(temp_dir, f"page_{i}_window.pdf")
        with open(chunk_file, "wb") as f:
            writer_local.write(f)
        with open(chunk_file, "rb") as f:
            pdf_bytes = f.read()
            
        prompt = f"""You are an expert RFI questionnaire extraction assistant.
        Review page {i} of the provided PDF snippet (which contains surrounding context from pages {window_pages} if applicable).
        Identify ONLY the RFI Questions or Questionnaire items that **BEGIN** on Page {i}. Do not extract questions that began on Page {i-1} (they will be captured by the previous page's run).
        You should use the surrounding context pages in the snippet to ensure you capture the **entire question text** even if it spans across into Page {i+1}.
        Return a list of questions found that begin on Page {i}.
        """
        
        response = call_llm_with_retry(
            client,
            model='gemini-2.5-pro',
            contents=[
                types.Part.from_bytes(data=pdf_bytes, mime_type='application/pdf'),
                prompt
            ],
            config={
                'response_mime_type': 'application/json',
                'response_schema': ExtractedQuestionsList,
            },
        )
        
        page_questions = []
        result = response.parsed
        if result and hasattr(result, 'questions'):
            for q in result.questions:
                page_questions.append({
                    "id": f"pdf_p{i}_{q.index_id}",
                    "text": q.question_text,
                    "location": {
                        "type": "pdf_page",
                        "page_number": i
                    }
                })
                
        os.remove(chunk_file)
        return page_questions

    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = [executor.submit(process_page, i) for i in range(1, total_pages + 1)]
        for future in futures:
            extracted_questions.extend(future.result())
            
    return extracted_questions
